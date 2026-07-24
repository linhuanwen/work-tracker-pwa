"""
工作清单 PWA - 透明窗口桌面启动器

后台 HTTP 服务器 + 无边框透明 WebView2 窗口 + 系统托盘。
HTML 透明区域可透出桌面壁纸，卡片以玻璃态悬浮。

原理：
  pywebview transparent=True → WinForms TransparencyKey = #010101
  → CSS html background: transparent → WebView2 不渲染背景
  → Form 背景 #010101 透出 → OS 级透明 → 桌面壁纸可见
  → body::before 半透明暖白叠加层提供小黄条式便签质感

窗口行为：
  · 启动时窗口自动停靠屏幕右边缘、最小宽度、全高（小黄条式侧面条）
  · 标题栏 X / Alt+F4 = 隐藏到托盘（进程与 HTTP 服务继续运行）
  · 托盘图标左键 = 显示主界面；右键菜单「退出」= 真正结束程序
  · 重复启动 = 通知已有实例显示窗口（单实例）

用法:
    pythonw launcher.py          （无控制台）
    python launcher.py --debug   （带控制台信息）
"""

import http.server
import os
import sys
import socket
import threading
import time
import json
import ctypes
import ctypes.wintypes
from pathlib import Path

PORT = 5173
HOST = "127.0.0.1"
URL = f"http://{HOST}:{PORT}"

# Transparency key — near-black, rare in UI, reliable for OS transparency
TKEY_HEX = '#010101'
TKEY_RGB = (1, 1, 1)

WINDOW_WIDTH = 520
WINDOW_HEIGHT = 780
WINDOW_TITLE = '工作清单'
WINDOW_MIN_WIDTH = 360    # logical (DPI-independent) pixels
WINDOW_MIN_HEIGHT = 480   # logical (DPI-independent) pixels


# ---------------------------------------------------------------------------
# AI Polish bridge — delegates to scripts/polish.py
# ---------------------------------------------------------------------------

def _get_scripts_dir():
    if getattr(sys, "frozen", False):
        return os.path.join(sys._MEIPASS, "scripts")
    else:
        return os.path.join(os.path.dirname(os.path.abspath(__file__)), "scripts")


def _resolve_ai_config(override: dict | None) -> dict:
    """Resolve AI config: use the request-supplied override when it carries an
    api_key, otherwise fall back to scripts/.env."""
    if override and str(override.get("api_key", "")).strip():
        return {
            "api_key": str(override.get("api_key", "")).strip(),
            "endpoint": str(override.get("endpoint", "")).strip()
            or "https://api.deepseek.com",
            "model": str(override.get("model", "")).strip() or "deepseek-chat",
        }

    scripts_dir = _get_scripts_dir()
    env_path = os.path.join(scripts_dir, ".env")
    from polish import load_config
    try:
        return load_config(env_path)
    except FileNotFoundError:
        raise Exception(
            "未配置 AI API。请在「设置 → AI 配置」中填写 API Key，"
            "或在 scripts/ 下创建 .env 并设置 AI_API_KEY。"
        )
    except ValueError as e:
        raise Exception(f"配置错误: {e}")


def _import_polish():
    scripts_dir = _get_scripts_dir()
    if scripts_dir not in sys.path:
        sys.path.insert(0, scripts_dir)
    try:
        from polish import build_polish_prompt, call_ai_api
    except ImportError:
        raise Exception(
            "润色脚本未找到，请确保 scripts/polish.py 和 scripts/.env 文件存在。"
        )
    return build_polish_prompt, call_ai_api


def _run_polish(raw_text: str, archive_type: str, config_override: dict | None = None) -> str:
    """Call polish.py functions to polish *raw_text* via AI.

    Returns the polished text, or raises an exception with a user-friendly
    message explaining what went wrong.
    """
    build_polish_prompt, call_ai_api = _import_polish()
    config = _resolve_ai_config(config_override)
    prompt = build_polish_prompt(raw_text)
    return call_ai_api(prompt, config)


# ---------------------------------------------------------------------------
# AI-generated summary Word document
# ---------------------------------------------------------------------------

def _summary_prompt(period_type: str, period_label: str, sections: dict) -> str:
    """Build a prompt asking the AI to produce a formal summary document."""
    section_lines = []
    for title, text in sections.items():
        section_lines.append(f"## {title}\n{text}\n")
    body = "\n".join(section_lines)

    return f"""你是一位资深文书助理。请根据以下{period_label}的工作材料，生成一份正式、简洁、结构化的工作总结 Word 文档内容。

## 写作要求
1. **文风**：正式、简洁、符合正式工作报告规范。
2. **用数据说话**：保留并突出量化产出和具体数据。
3. **避免口语化**：删除"搞定了""推进了一下"等日常表达。
4. **避免情绪化**：不添加"极大地""非常"等主观修饰词。
5. **结构化**：使用 Markdown 标题（# 一级标题、## 二级标题）和项目符号列表组织内容。
6. **不编造**：不增加原文没有的信息，不删除原文已有的事实。

## 输出格式
只输出 Markdown 格式的文档正文，不要添加任何解释、标记或前缀。第一行应为一级标题，例如"{period_label}工作总结"。

## 原始材料
{body}

请开始生成："""


def _markdown_to_docx(markdown_text: str, output_path: str) -> None:
    """Convert simple Markdown to a .docx file (Chinese-friendly)."""
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import re

    doc = Document()
    # Set default font for the document
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(10.5)
    # Set narrow margins to make better use of the narrow window if printed
    sections = doc.sections[0]
    sections.top_margin = Inches(0.6)
    sections.bottom_margin = Inches(0.6)
    sections.left_margin = Inches(0.7)
    sections.right_margin = Inches(0.7)

    for raw_line in markdown_text.splitlines():
        line = raw_line.rstrip()
        if not line.strip():
            continue
        stripped = line.lstrip()
        if stripped.startswith('# '):
            p = doc.add_heading(stripped[2:], level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif stripped.startswith('## '):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith('### '):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith('- ') or stripped.startswith('* '):
            doc.add_paragraph(stripped[2:], style='List Bullet')
        elif re.match(r'^\d+\.\s', stripped):
            text = re.sub(r'^\d+\.\s', '', stripped)
            doc.add_paragraph(text, style='List Number')
        else:
            doc.add_paragraph(line)

    doc.save(output_path)


def _type_label_map() -> dict[str, tuple[str, str, str]]:
    """Map summary type to (folder_name, doc_name_prefix, period_label)."""
    return {
        'week': ('周报', '周', '周'),
        'month': ('月报', '月', '月'),
        'year': ('年报', '年度', '年'),
    }


def _summary_doc_filename(period_type: str, key: str) -> tuple[str, str]:
    """Return (folder_name, filename.docx) for the given type/key."""
    folder, _, _ = _type_label_map()[period_type]
    if period_type == 'week':
        # key format: "2026-W29"
        year, week = key.split('-W')
        filename = f"{year}年第{int(week)}周工作总结.docx"
    elif period_type == 'month':
        # key format: "2026-07"
        year, month = key.split('-')
        filename = f"{year}年{int(month)}月工作总结.docx"
    else:
        # key format: "2026"
        filename = f"{key}年度工作总结.docx"
    return folder, filename


def _generate_summary_doc(period_type: str, key: str, sections: dict, config_override: dict | None = None) -> str:
    """Generate a Word document summary via AI and save it to the shared folder.

    Returns the saved file path. Raises an exception on failure.
    """
    if not _data_folder_path:
        raise Exception("未配置数据文件夹，无法保存总结文档。")

    scripts_dir = _get_scripts_dir()
    if scripts_dir not in sys.path:
        sys.path.insert(0, scripts_dir)

    try:
        from polish import call_ai_api
    except ImportError:
        raise Exception("AI 润色脚本未找到，请确保 scripts/polish.py 存在。")

    config = _resolve_ai_config(config_override)

    _, period_word, _ = _type_label_map()[period_type]
    if period_type == 'week':
        year, week = key.split('-W')
        full_label = f"{year}年第{int(week)}周"
    elif period_type == 'month':
        year, month = key.split('-')
        full_label = f"{year}年{int(month)}月"
    else:
        full_label = f"{key}年度"

    prompt = _summary_prompt(period_type, full_label, sections)
    markdown = call_ai_api(prompt, config)

    folder_name, filename = _summary_doc_filename(period_type, key)
    output_dir = Path(_data_folder_path) / folder_name
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / filename

    _markdown_to_docx(markdown, str(output_path))
    return str(output_path)


# ---- Win32 API for window resize & drag ----
user32 = ctypes.windll.user32
user32.FindWindowW.argtypes = [ctypes.wintypes.LPCWSTR, ctypes.wintypes.LPCWSTR]
user32.FindWindowW.restype = ctypes.wintypes.HWND
user32.SetWindowPos.argtypes = [
    ctypes.wintypes.HWND, ctypes.wintypes.HWND,
    ctypes.c_int, ctypes.c_int, ctypes.c_int, ctypes.c_int,
    ctypes.wintypes.UINT,
]
user32.SetWindowPos.restype = ctypes.wintypes.BOOL
user32.GetWindowRect.argtypes = [ctypes.wintypes.HWND, ctypes.c_void_p]
user32.GetWindowRect.restype = ctypes.wintypes.BOOL
user32.ReleaseCapture.argtypes = []
user32.ReleaseCapture.restype = ctypes.wintypes.BOOL
user32.SendMessageW.argtypes = [
    ctypes.wintypes.HWND, ctypes.wintypes.UINT,
    ctypes.wintypes.WPARAM, ctypes.wintypes.LPARAM,
]
user32.SendMessageW.restype = ctypes.wintypes.LPARAM
user32.PostMessageW.argtypes = [
    ctypes.wintypes.HWND, ctypes.wintypes.UINT,
    ctypes.wintypes.WPARAM, ctypes.wintypes.LPARAM,
]
user32.PostMessageW.restype = ctypes.wintypes.BOOL
user32.GetCursorPos.argtypes = [ctypes.c_void_p]
user32.GetCursorPos.restype = ctypes.wintypes.BOOL
user32.ShowWindow.argtypes = [ctypes.wintypes.HWND, ctypes.c_int]
user32.ShowWindow.restype = ctypes.wintypes.BOOL
user32.IsZoomed.argtypes = [ctypes.wintypes.HWND]
user32.IsZoomed.restype = ctypes.wintypes.BOOL
user32.GetDpiForWindow.argtypes = [ctypes.wintypes.HWND]
user32.GetDpiForWindow.restype = ctypes.wintypes.UINT
user32.SetForegroundWindow.argtypes = [ctypes.wintypes.HWND]
user32.SetForegroundWindow.restype = ctypes.wintypes.BOOL
user32.SystemParametersInfoW.argtypes = [
    ctypes.wintypes.UINT, ctypes.wintypes.UINT,
    ctypes.c_void_p, ctypes.wintypes.UINT,
]
user32.SystemParametersInfoW.restype = ctypes.wintypes.BOOL

SWP_NOZORDER = 0x0004
SWP_NOMOVE = 0x0002
SWP_NOACTIVATE = 0x0010

SW_HIDE = 0
SW_MINIMIZE = 6
SW_MAXIMIZE = 3
SW_RESTORE = 9

SPI_GETWORKAREA = 0x0030

WM_NCLBUTTONDOWN = 0x00A1
HTCAPTION = 2


class POINT(ctypes.Structure):
    _fields_ = [("x", ctypes.c_long), ("y", ctypes.c_long)]


class RECT(ctypes.Structure):
    _fields_ = [
        ("left", ctypes.c_long),
        ("top", ctypes.c_long),
        ("right", ctypes.c_long),
        ("bottom", ctypes.c_long),
    ]


# Global: window HWND (populated after window appears)
_window_hwnd = None
_hwnd_ready = threading.Event()
_debug_mode = False

# Tray / quit state
_webview_window = None   # pywebview Window (kept for events + destroy)
_tray_icon = None        # pystray.Icon
_allow_quit = False      # tray「退出」置 True 后才允许真正关窗
_docked = False          # 启动停靠只执行一次

# NOTE: frameless drag & resize are driven by the frontend over the
# /api/window HTTP bridge.  A previous implementation subclassed the
# top-level WNDPROC to intercept WM_NCHITTEST, but the WebView2 child
# window covers the entire client area, so the top-level window never
# receives hit-test messages — edge/corner resize zones never fired.
# The JS bridge (start_drag / move_resize / minimize / toggle_maximize /
# close) works regardless of child windows.


def find_and_store_hwnd():
    """Daemon: poll until the window HWND is found and cache it for API use."""
    global _window_hwnd
    for _ in range(100):  # ~10 seconds
        hwnd = user32.FindWindowW(None, WINDOW_TITLE)
        if hwnd:
            _window_hwnd = hwnd
            if _debug_mode:
                print(f"[*] 窗口句柄已捕获 (HWND={hwnd})")
            _hwnd_ready.set()
            return
        time.sleep(0.1)

    if _debug_mode:
        print("[!] 未能在 10 秒内找到窗口句柄，窗口控制 API 可能失效")


def resize_window(width: int, height: int) -> bool:
    """Resize the window to the given dimensions."""
    hwnd = _window_hwnd
    if not hwnd:
        hwnd = user32.FindWindowW(None, WINDOW_TITLE)
    if not hwnd:
        return False
    return user32.SetWindowPos(
        hwnd, 0, 0, 0, width, height,
        SWP_NOZORDER | SWP_NOMOVE | SWP_NOACTIVATE,
    ) != 0


def move_resize_window(left: int, top: int, width: int, height: int) -> bool:
    """Move and resize the window to the given position and dimensions."""
    hwnd = _window_hwnd
    if not hwnd:
        hwnd = user32.FindWindowW(None, WINDOW_TITLE)
    if not hwnd:
        return False
    return user32.SetWindowPos(
        hwnd, 0, left, top, width, height,
        SWP_NOZORDER | SWP_NOACTIVATE,
    ) != 0


def start_drag_window() -> bool:
    """Start a native drag operation from the title-bar area.

    Sends WM_NCLBUTTONDOWN with HTCAPTION so Windows runs its modal move
    loop.  lParam must carry the *actual* cursor position in screen
    coordinates — passing 0 made the move loop anchor at screen (0,0)
    and the window jumped on every drag.
    """
    hwnd = _window_hwnd
    if not hwnd:
        hwnd = user32.FindWindowW(None, WINDOW_TITLE)
    if not hwnd:
        return False
    pt = POINT()
    if not user32.GetCursorPos(ctypes.byref(pt)):
        return False
    lparam = ((pt.y & 0xFFFF) << 16) | (pt.x & 0xFFFF)
    user32.ReleaseCapture()
    user32.SendMessageW(hwnd, WM_NCLBUTTONDOWN, HTCAPTION, lparam)
    return True


def minimize_window() -> bool:
    """Minimize the window to the taskbar."""
    hwnd = _window_hwnd
    if not hwnd:
        hwnd = user32.FindWindowW(None, WINDOW_TITLE)
    if not hwnd:
        return False
    user32.ShowWindow(hwnd, SW_MINIMIZE)
    return True


def toggle_maximize_window() -> tuple[bool, bool]:
    """Toggle maximize/restore. Returns (ok, maximized_after)."""
    hwnd = _window_hwnd
    if not hwnd:
        hwnd = user32.FindWindowW(None, WINDOW_TITLE)
    if not hwnd:
        return False, False
    if user32.IsZoomed(hwnd):
        user32.ShowWindow(hwnd, SW_RESTORE)
        return True, False
    user32.ShowWindow(hwnd, SW_MAXIMIZE)
    return True, True


def is_window_maximized() -> bool:
    """Return whether the window is currently maximized."""
    hwnd = _window_hwnd
    if not hwnd:
        hwnd = user32.FindWindowW(None, WINDOW_TITLE)
    if not hwnd:
        return False
    return bool(user32.IsZoomed(hwnd))


def get_min_window_size() -> tuple[int, int]:
    """OS-enforced minimum size in *physical* pixels.

    pywebview's min_size is logical (DPI-independent); WinForms scales it
    by the window DPI (e.g. 480 -> 720 physical at 150% scaling).  The
    frontend clamps resize in physical pixels, so it must use the same
    values the OS enforces — otherwise the anchored edge slides once the
    OS clamp kicks in.
    """
    hwnd = _window_hwnd or user32.FindWindowW(None, WINDOW_TITLE)
    scale = 1.0
    if hwnd:
        try:
            dpi = user32.GetDpiForWindow(hwnd)
            if dpi:
                scale = dpi / 96.0
        except Exception:
            pass
    return (round(WINDOW_MIN_WIDTH * scale), round(WINDOW_MIN_HEIGHT * scale))


def dock_right():
    """启动布局：窗口停靠到主屏工作区右边缘、全高、最小宽度（小黄条式侧面条）。

    SystemParametersInfoW 返回物理像素（进程 DPI 感知），与 SetWindowPos
    单位一致。"""
    global _docked
    if _docked:
        return
    hwnd = _window_hwnd
    if not hwnd:
        return
    wa = RECT()
    if not user32.SystemParametersInfoW(SPI_GETWORKAREA, 0, ctypes.byref(wa), 0):
        return
    min_w, _ = get_min_window_size()
    w = min_w
    h = wa.bottom - wa.top
    if move_resize_window(wa.right - w, wa.top, w, h):
        _docked = True


def hide_window() -> bool:
    """隐藏窗口到托盘（屏幕和任务栏都不见，进程与 HTTP 服务继续跑）。"""
    hwnd = _window_hwnd or user32.FindWindowW(None, WINDOW_TITLE)
    if not hwnd:
        return False
    user32.ShowWindow(hwnd, SW_HIDE)
    return True


def show_window() -> bool:
    """从托盘恢复窗口并置前。"""
    hwnd = _window_hwnd or user32.FindWindowW(None, WINDOW_TITLE)
    if not hwnd:
        return False
    user32.ShowWindow(hwnd, SW_RESTORE)  # 同时解除最小化/隐藏
    user32.SetForegroundWindow(hwnd)
    return True


def make_tray_image():
    """内存中绘制托盘图标（无需外部图标文件/字体）。"""
    from PIL import Image, ImageDraw

    img = Image.new("RGBA", (64, 64), (0, 0, 0, 0))
    d = ImageDraw.Draw(img)
    d.rounded_rectangle([2, 2, 62, 62], radius=14, fill=(74, 108, 247, 255))
    white = (255, 255, 255, 255)
    # 第一行：勾 + 横线；下面两行：圆点 + 横线 —— 缩小后仍可读作"清单"
    d.line([(14, 20), (19, 26), (28, 13)], fill=white, width=5, joint="curve")
    d.line([(34, 20), (50, 20)], fill=white, width=5)
    d.ellipse([17, 30, 23, 36], fill=white)
    d.line([(34, 33), (50, 33)], fill=white, width=5)
    d.ellipse([17, 43, 23, 49], fill=white)
    d.line([(34, 46), (50, 46)], fill=white, width=5)
    return img


def start_tray():
    """启动系统托盘图标（左键=显示主界面，右键菜单含退出）。

    pystray 的 win32 后端在自己的线程里跑消息循环；本函数在 daemon
    线程中调用 icon.run()。托盘「退出」是唯一真正结束程序的途径。
    """
    global _tray_icon
    import pystray

    menu = pystray.Menu(
        pystray.MenuItem("显示主界面", lambda icon, item: show_window(), default=True),
        pystray.MenuItem("退出", quit_app),
    )
    _tray_icon = pystray.Icon("wjl-tray", make_tray_image(), WINDOW_TITLE, menu)
    threading.Thread(target=_tray_icon.run, daemon=True).start()


def on_window_closing():
    """pywebview closing 事件：默认否决关闭、改为隐藏到托盘。"""
    if _allow_quit:
        return True
    hide_window()
    return False


def quit_app(icon=None, item=None):
    """托盘菜单「退出」：真正结束程序。"""
    global _allow_quit
    _allow_quit = True
    if _debug_mode:
        print("[*] 托盘退出，正在关闭…", flush=True)
    if _tray_icon is not None:
        try:
            _tray_icon.stop()
        except Exception:
            pass
    win = _webview_window
    if win is not None:
        win.destroy()  # closing 事件因 _allow_quit 放行 → webview.start() 返回
    else:
        os._exit(0)    # 窗口还没建成就收到退出（罕见竞态）


def notify_existing_instance() -> bool:
    """向已在运行的实例 POST show。收到任何 HTTP 响应都视为成功。"""
    import urllib.error
    import urllib.request

    req = urllib.request.Request(
        f"{URL}/api/window",
        data=b'{"action":"show"}',
        headers={"Content-Type": "application/json"},
    )
    try:
        urllib.request.urlopen(req, timeout=2)
        return True
    except urllib.error.HTTPError:
        return True   # 有响应（如老实例不认识 show）—— 一样交给旧实例
    except Exception:
        return False  # 连接失败：旧实例已死，继续正常启动


def get_window_rect():
    """Get current window rect as (left, top, right, bottom)."""
    hwnd = _window_hwnd
    if not hwnd:
        hwnd = user32.FindWindowW(None, WINDOW_TITLE)
    if not hwnd:
        return None

    rect = RECT()
    if user32.GetWindowRect(hwnd, ctypes.byref(rect)):
        return (rect.left, rect.top, rect.right, rect.bottom)
    return None


# ---------------------------------------------------------------------------
# State persistence (.wjl-state.json)
# ---------------------------------------------------------------------------

def get_state_path():
    if getattr(sys, "frozen", False):
        base = os.path.dirname(sys.executable)
    else:
        base = os.path.dirname(os.path.abspath(__file__))
    return os.path.join(base, ".wjl-state.json")


def read_state():
    try:
        with open(get_state_path(), "r", encoding="utf-8") as f:
            return json.load(f)
    except (FileNotFoundError, json.JSONDecodeError):
        return {}


def write_state(state):
    with open(get_state_path(), "w", encoding="utf-8") as f:
        json.dump(state, f, ensure_ascii=False, indent=2)


# ---------------------------------------------------------------------------
# Data folder configuration
# ---------------------------------------------------------------------------

_data_folder_path: str | None = None


def _get_search_bases() -> list[str]:
    """Return directories to search for wjl-config.txt / data.json."""
    bases: list[str] = []
    # 1. Directory containing the exe / launcher.py
    if getattr(sys, "frozen", False):
        bases.append(os.path.dirname(sys.executable))
    else:
        bases.append(os.path.dirname(os.path.abspath(__file__)))
    # 2. Current working directory
    bases.append(os.getcwd())
    return bases


def _try_read_config_txt(debug: bool = False) -> str | None:
    """Read the first valid directory path from any wjl-config.txt found.

    Supports:
      - UTF-8 with or without BOM
      - Absolute paths
      - Relative paths (resolved relative to the config file's directory)
    """
    for base in _get_search_bases():
        config_file = os.path.join(base, "wjl-config.txt")
        if os.path.isfile(config_file):
            try:
                raw = Path(config_file).read_text("utf-8")
                path = raw.strip().lstrip("\ufeff")
                if not path:
                    continue
                if not os.path.isabs(path):
                    path = os.path.join(os.path.dirname(config_file), path)
                path = os.path.abspath(path)
                if os.path.isdir(path):
                    if debug:
                        print(f"[*] 从 {config_file} 读取到数据文件夹: {path}")
                    return path
                if debug:
                    print(f"[!] {config_file} 中的路径不存在或不是目录: {path}")
            except Exception as e:
                if debug:
                    print(f"[!] 读取 {config_file} 失败: {e}")
                continue
    return None


def _try_find_data_json() -> str | None:
    """If a data.json exists next to the exe or in cwd, use its parent folder."""
    for base in _get_search_bases():
        candidate = os.path.join(base, "data.json")
        if os.path.isfile(candidate):
            return os.path.abspath(base)
    return None


def get_configured_data_folder(debug: bool = False) -> str | None:
    """Return the configured data folder path, or None if not configured.

    Precedence:
      1. wjl-config.txt in launcher / cwd (single line, UTF-8)
      2. dataFolderPath key in .wjl-state.json
      3. A data.json file found next to the exe / in cwd
    """
    # 1. Explicit config file
    path = _try_read_config_txt(debug=debug)
    if path:
        return path

    # 2. State file persisted by launcher/PWA
    state = read_state()
    path = state.get("dataFolderPath")
    if path and os.path.isdir(path):
        if debug:
            print(f"[*] 从 .wjl-state.json 读取到数据文件夹: {path}")
        return os.path.abspath(path)

    # 3. Auto-detect existing data.json
    path = _try_find_data_json()
    if path:
        if debug:
            print(f"[*] 自动发现 data.json 所在文件夹: {path}")
        return path

    if debug:
        print("[!] 未找到配置的数据文件夹")
    return None


def _default_data_dict() -> dict:
    from datetime import datetime, timezone
    return {
        "version": 1,
        "lastModified": datetime.now(timezone.utc).isoformat(),
        "settings": {
            "weeklySummaryDay": 5,
            "monthlySummaryDay": 28,
            "aiPolishFlag": False,
            "categories": [
                "人员调配", "内部招聘", "奖惩管理", "绩效管理",
                "劳动关系", "领导交办", "其他",
            ],
        },
        "projects": [],
        "tasks": [],
        "archives": {"weeks": {}, "months": {}, "years": {}},
    }


def read_data_json() -> dict | None:
    """Read data.json from the configured data folder."""
    if not _data_folder_path:
        return None
    data_path = Path(_data_folder_path) / "data.json"
    if not data_path.exists():
        default_data = _default_data_dict()
        write_data_json(default_data)
        return default_data
    try:
        with open(data_path, "r", encoding="utf-8") as f:
            return json.load(f)
    except (FileNotFoundError, json.JSONDecodeError):
        return _default_data_dict()


def write_data_json(data: dict) -> bool:
    """Write data.json to the configured data folder atomically."""
    if not _data_folder_path:
        return False
    data_path = Path(_data_folder_path) / "data.json"
    try:
        data_path.parent.mkdir(parents=True, exist_ok=True)
        tmp_path = data_path.with_suffix(".tmp")
        with open(tmp_path, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        tmp_path.replace(data_path)
        return True
    except Exception:
        return False


# ---------------------------------------------------------------------------
# HTTP Server
# ---------------------------------------------------------------------------

def get_dist_dir():
    if getattr(sys, "frozen", False):
        bundle_dir = sys._MEIPASS
    else:
        bundle_dir = os.path.dirname(os.path.abspath(__file__))

    for candidate in [
        os.path.join(bundle_dir, "dist"),
        os.path.join(os.path.dirname(os.path.abspath(__file__)), "dist"),
        os.path.join(os.getcwd(), "dist"),
    ]:
        if os.path.isdir(candidate):
            return candidate
    sys.exit("ERROR: dist/ directory not found")


class RequestHandler(http.server.SimpleHTTPRequestHandler):
    def do_GET(self):
        if self.path == '/api/state':
            self.send_response(200)
            self.send_header('Content-type', 'application/json; charset=utf-8')
            self.send_header('Cache-Control', 'no-store')
            self.end_headers()
            state = read_state()
            if _data_folder_path:
                state["dataFolderPath"] = _data_folder_path
            self.wfile.write(json.dumps(state, ensure_ascii=False).encode('utf-8'))
            return

        if self.path == '/api/data':
            self.send_response(200)
            self.send_header('Content-type', 'application/json; charset=utf-8')
            self.send_header('Cache-Control', 'no-store')
            self.end_headers()
            data = read_data_json()
            if data is None:
                self.send_response(503)
                self.end_headers()
                self.wfile.write(b'{"error":"data folder not configured"}')
                return
            self.wfile.write(json.dumps(data, ensure_ascii=False).encode('utf-8'))
            return

        if self.path == '/api/window':
            # 首个 GET 时窗口创建与 DPI 缩放已彻底完成 —— 在此做一次性
            # 启动停靠（内部有 _docked 守卫，后续 GET 直接返回）
            dock_right()
            rect = get_window_rect()
            maximized = is_window_maximized()
            min_w, min_h = get_min_window_size()
            self.send_response(200)
            self.send_header('Content-type', 'application/json; charset=utf-8')
            self.end_headers()
            if rect:
                self.wfile.write(json.dumps({
                    "left": rect[0], "top": rect[1],
                    "right": rect[2], "bottom": rect[3],
                    "width": rect[2] - rect[0],
                    "height": rect[3] - rect[1],
                    "maximized": maximized,
                    "minWidth": min_w,
                    "minHeight": min_h,
                }).encode('utf-8'))
            else:
                self.wfile.write(json.dumps({
                    "width": 520, "height": 780, "maximized": False,
                    "minWidth": min_w, "minHeight": min_h,
                }).encode('utf-8'))
            return

        if self.path == '/shutdown':
            self.send_response(200)
            self.send_header('Content-type', 'text/plain; charset=utf-8')
            self.end_headers()
            self.wfile.write(b'OK')
            threading.Thread(target=self.server.shutdown, daemon=True).start()
            return

        super().do_GET()

    def do_POST(self):
        if self.path == '/api/state':
            length = int(self.headers.get('Content-Length', 0))
            body = self.rfile.read(length)
            try:
                state = json.loads(body)
                write_state(state)
                self.send_response(200)
                self.send_header('Content-type', 'application/json; charset=utf-8')
                self.end_headers()
                self.wfile.write(b'{"ok":true}')
            except json.JSONDecodeError:
                self.send_response(400)
                self.end_headers()
            return

        if self.path == '/api/window':
            length = int(self.headers.get('Content-Length', 0))
            body = self.rfile.read(length)
            try:
                data = json.loads(body)
                action = data.get("action", "")
                if _debug_mode:
                    print(f"[*] /api/window POST: {data}", flush=True)
                if action == "resize":
                    w = int(data.get("width", WINDOW_WIDTH))
                    h = int(data.get("height", WINDOW_HEIGHT))
                    ok = resize_window(w, h)
                    self.send_response(200)
                    self.send_header('Content-type', 'application/json; charset=utf-8')
                    self.end_headers()
                    self.wfile.write(json.dumps({"ok": ok}).encode('utf-8'))
                    return
                if action == "move_resize":
                    x = int(data.get("x", 0))
                    y = int(data.get("y", 0))
                    w = int(data.get("width", WINDOW_WIDTH))
                    h = int(data.get("height", WINDOW_HEIGHT))
                    ok = move_resize_window(x, y, w, h)
                    self.send_response(200)
                    self.send_header('Content-type', 'application/json; charset=utf-8')
                    self.end_headers()
                    self.wfile.write(json.dumps({"ok": ok}).encode('utf-8'))
                    return
                if action == "start_drag":
                    ok = start_drag_window()
                    self.send_response(200)
                    self.send_header('Content-type', 'application/json; charset=utf-8')
                    self.end_headers()
                    self.wfile.write(json.dumps({"ok": ok}).encode('utf-8'))
                    return
                if action == "minimize":
                    ok = minimize_window()
                    self.send_response(200)
                    self.send_header('Content-type', 'application/json; charset=utf-8')
                    self.end_headers()
                    self.wfile.write(json.dumps({"ok": ok}).encode('utf-8'))
                    return
                if action == "toggle_maximize":
                    ok, maximized = toggle_maximize_window()
                    self.send_response(200)
                    self.send_header('Content-type', 'application/json; charset=utf-8')
                    self.end_headers()
                    self.wfile.write(json.dumps({"ok": ok, "maximized": maximized}).encode('utf-8'))
                    return
                if action == "close":
                    # 关闭按钮 = 隐藏到托盘；真正退出只走托盘菜单「退出」
                    ok = hide_window()
                    self.send_response(200)
                    self.send_header('Content-type', 'application/json; charset=utf-8')
                    self.end_headers()
                    self.wfile.write(json.dumps({"ok": ok}).encode('utf-8'))
                    return
                if action == "show":
                    ok = show_window()
                    self.send_response(200)
                    self.send_header('Content-type', 'application/json; charset=utf-8')
                    self.end_headers()
                    self.wfile.write(json.dumps({"ok": ok}).encode('utf-8'))
                    return
            except (json.JSONDecodeError, ValueError):
                pass
            self.send_response(400)
            self.end_headers()
            return

        if self.path == '/api/data':
            length = int(self.headers.get('Content-Length', 0))
            body = self.rfile.read(length)
            try:
                payload = json.loads(body)
                ok = write_data_json(payload)
                self.send_response(200 if ok else 500)
                self.send_header('Content-type', 'application/json; charset=utf-8')
                self.end_headers()
                self.wfile.write(json.dumps({"ok": ok}).encode('utf-8'))
            except json.JSONDecodeError:
                self.send_response(400)
                self.end_headers()
            return

        if self.path == '/api/summary':
            length = int(self.headers.get('Content-Length', 0))
            body = self.rfile.read(length)
            try:
                payload = json.loads(body)
                period_type = payload.get('type', 'week')
                key = payload.get('key', '')
                sections = payload.get('sections', {})
                if period_type not in ('week', 'month', 'year'):
                    self.send_response(400)
                    self.send_header('Content-type', 'application/json; charset=utf-8')
                    self.end_headers()
                    self.wfile.write(json.dumps({"ok": False, "error": "invalid type"}).encode('utf-8'))
                    return
                if not key:
                    self.send_response(400)
                    self.send_header('Content-type', 'application/json; charset=utf-8')
                    self.end_headers()
                    self.wfile.write(json.dumps({"ok": False, "error": "key is empty"}).encode('utf-8'))
                    return

                saved_path = _generate_summary_doc(period_type, key, sections, payload.get('config'))
                self.send_response(200)
                self.send_header('Content-type', 'application/json; charset=utf-8')
                self.end_headers()
                self.wfile.write(json.dumps({"ok": True, "path": saved_path}).encode('utf-8'))
            except Exception as e:
                self.send_response(200)
                self.send_header('Content-type', 'application/json; charset=utf-8')
                self.end_headers()
                self.wfile.write(json.dumps({"ok": False, "error": str(e)}).encode('utf-8'))
            return

        if self.path == '/api/polish':
            length = int(self.headers.get('Content-Length', 0))
            body = self.rfile.read(length)
            try:
                payload = json.loads(body)
                raw_text = payload.get('text', '')
                archive_type = payload.get('type', 'week')
                if not raw_text or not raw_text.strip():
                    self.send_response(400)
                    self.send_header('Content-type', 'application/json; charset=utf-8')
                    self.end_headers()
                    self.wfile.write(json.dumps({"ok": False, "error": "text is empty"}).encode('utf-8'))
                    return

                polished = _run_polish(raw_text, archive_type, payload.get('config'))
                self.send_response(200)
                self.send_header('Content-type', 'application/json; charset=utf-8')
                self.end_headers()
                self.wfile.write(json.dumps({"ok": True, "polished": polished}).encode('utf-8'))
            except Exception as e:
                self.send_response(200)  # Return 200 with error flag — avoid breaking the PWA
                self.send_header('Content-type', 'application/json; charset=utf-8')
                self.end_headers()
                self.wfile.write(json.dumps({"ok": False, "error": str(e)}).encode('utf-8'))
            return

        self.send_response(404)
        self.end_headers()

    def log_message(self, format, *args):
        pass


def start_server():
    dist_dir = get_dist_dir()
    os.chdir(dist_dir)
    server = http.server.ThreadingHTTPServer((HOST, PORT), RequestHandler)
    t = threading.Thread(target=server.serve_forever, daemon=True)
    t.start()
    return server


def is_server_running():
    try:
        s = socket.create_connection((HOST, PORT), timeout=1)
        s.close()
        return True
    except (ConnectionRefusedError, OSError):
        return False


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    debug = "--debug" in sys.argv
    global _debug_mode, _data_folder_path, _webview_window
    _debug_mode = debug

    # Load configured data folder (if any) so /api/data can serve it
    _data_folder_path = get_configured_data_folder(debug=debug)
    if debug:
        if _data_folder_path:
            print(f"[*] 数据文件夹: {_data_folder_path}")
        else:
            print("[!] 没有配置数据文件夹；启动后将要求手动选择文件夹")

    # 0. 单实例：已有实例在跑 → 通知它显示窗口，本进程直接退出
    #    （避免第二个窗口 + 第二个托盘图标）
    if is_server_running():
        if notify_existing_instance():
            if debug:
                print("[*] 已有实例在运行，已通知其显示窗口，本进程退出")
            sys.exit(0)
        if debug:
            print("[!] 端口被占但旧实例无响应，继续启动")

    # 1. Start HTTP server
    start_server()
    for _ in range(20):
        if is_server_running():
            break
        time.sleep(0.05)
    if debug:
        print(f"Server started at {URL}")

    # 2. Create window
    import webview

    _webview_window = webview.create_window(
        title=WINDOW_TITLE,
        url=URL,
        width=WINDOW_WIDTH,
        height=WINDOW_HEIGHT,
        resizable=True,
        # OS 最小尺寸只是地板：收起模式（36px 标题栏）必须低于 UI 最小尺寸
        # （360×480 由 get_min_window_size 上报，前端在展开态自行钳制）
        min_size=(240, 32),
        frameless=True,
        transparent=True,
        background_color=TKEY_HEX,
        on_top=False,
        confirm_close=False,
        text_select=True,
        # easy_drag (default True) makes pywebview track EVERY mousedown in
        # the page and call its own window.move() on mousemove — which fights
        # our /api/window drag & resize bridge and throws NoneType errors in
        # SetWindowPos.  Our TitleBar/resize handles replace it entirely.
        easy_drag=False,
    )
    # 拦截真正的窗口关闭（标题栏 X / Alt+F4 等）：默认改为隐藏到托盘，
    # 只有托盘菜单「退出」（_allow_quit=True）才放行
    _webview_window.events.closing += on_window_closing

    # 3. 系统托盘（左键=显示主界面，右键菜单含「退出」）
    start_tray()

    # 4. Background thread to capture HWND for API use
    hwnd_thread = threading.Thread(target=find_and_store_hwnd, daemon=True)
    hwnd_thread.start()

    # 5. GUI message loop (blocks until closed)
    webview.start(gui='edgechromium')


if __name__ == "__main__":
    main()
